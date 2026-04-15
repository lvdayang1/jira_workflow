# -*- coding: utf-8 -*-
"""
测试用例文档生成器 - 根据 JSON 生成测试用例文档
支持自定义 Excel 模板
"""
import os
import json
import re
from pathlib import Path

# 默认的生成器字段顺序（当模板没有指定时使用）
DEFAULT_COLUMNS = [
    '用例编号', '用例标题', '测试模块', '测试类型', '优先级',
    '前置条件', '测试步骤', '预期结果', '测试数据', '备注'
]


class JSONTestCaseParser:
    """解析 JSON 格式的测试用例"""

    def __init__(self, json_path):
        self.json_path = json_path
        self.test_cases = []
        self.ticket_info = {}
        self._parse()

    def _parse(self):
        with open(self.json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        self._extract_ticket_info(data)
        self._extract_test_cases(data)

    def _extract_ticket_info(self, data):
        self.ticket_info = {
            "Ticket ID": data.get("id", ""),
            "标题": data.get("summary", ""),
            "优先级": data.get("priority", ""),
            "状态": data.get("status", ""),
            "报告人": data.get("reporter", ""),
            "指派人": data.get("assignee", "")
        }

    def _extract_test_cases(self, data):
        test_cases = data.get("test_cases", [])
        for idx, tc in enumerate(test_cases, 1):
            case_data = {
                'id': tc.get('id', f'TC-{idx:03d}'),
                'name': tc.get('name', tc.get('用例标题', '')),
                'module': tc.get('module', tc.get('测试模块', '')),
                'test_type': tc.get('type', tc.get('测试类型', '功能测试')),
                'priority': tc.get('priority', tc.get('优先级', '中')),
                'precondition': tc.get('precondition', tc.get('前提条件', '')),
                'steps': tc.get('steps', tc.get('测试步骤', [])),
                'expected_results': tc.get('expected_results', tc.get('预期结果', [])),
                'test_data': tc.get('test_data', tc.get('测试数据', {})),
                'remarks': tc.get('remarks', tc.get('备注', '')),
                # 支持用户模板字段
                'operation': tc.get('operation', tc.get('操作', '')),
                'input': tc.get('input', tc.get('输入', '')),
                'test_result': tc.get('test_result', tc.get('测试结果', ''))
            }
            # 处理 steps 和 expected_results 可能是字符串的情况
            if isinstance(case_data['steps'], str):
                case_data['steps'] = [s.strip() for s in case_data['steps'].split('\n') if s.strip()]
            if isinstance(case_data['expected_results'], str):
                case_data['expected_results'] = [s.strip() for s in case_data['expected_results'].split('\n') if s.strip()]
            if isinstance(case_data['test_data'], dict):
                case_data['test_data'] = json.dumps(case_data['test_data'], ensure_ascii=False)
            self.test_cases.append(case_data)


class MarkdownTestCaseParser:
    """解析 Markdown 格式的测试用例"""

    def __init__(self, md_content):
        self.md_content = md_content
        self.test_cases = []
        self.ticket_info = {}
        self._parse()

    def _parse(self):
        self._parse_ticket_info()
        self._parse_test_cases()

    def _clean_html_tags(self, text):
        text = re.sub(r'<br\s*/?>', '\n', text)
        text = re.sub(r'<[^>]+>', '', text)
        return text.strip()

    def _parse_table_content(self, table_text):
        lines = table_text.strip().split('\n')
        result = {}
        for line in lines:
            match = re.match(r'\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|', line)
            if match:
                key = match.group(1).strip()
                value = match.group(2).strip()
                value = self._clean_html_tags(value)
                result[key] = value
        return result

    def _parse_ticket_info(self):
        info_match = re.search(r'## 基本信息\s*\n([\s\S]*?)(?=\n## |\n---)', self.md_content)
        if info_match:
            info_text = info_match.group(1)
            self.ticket_info = self._parse_table_content(info_text)

    def _parse_test_cases(self):
        case_pattern = r'### 测试用例编号：(.+?)\n([\s\S]*?)(?=\n### 测试用例|\n## |\n---+\s*$)'
        matches = re.findall(case_pattern, self.md_content)

        for idx, match in enumerate(matches, 1):
            case_name = match[0].strip()
            case_content = match[1]

            case_data = {
                'id': f'TC-{idx:03d}',
                'name': case_name,
                'module': '',
                'test_type': '功能测试',
                'priority': '中',
                'precondition': '',
                'steps': [],
                'expected_results': [],
                'test_data': '',
                'remarks': ''
            }

            precondition_match = re.search(r'\*\*前置条件\*\*：\n([\s\S]*?)(?=\n\*\*测试步骤|$)', case_content, re.DOTALL)
            if precondition_match:
                case_data['precondition'] = precondition_match.group(1).strip()

            steps_lines = []
            step_pattern = r'(?:^|\n)(\d+)[.、]\s*(.+)'
            step_matches = re.findall(step_pattern, case_content)
            for step_num, step_text in step_matches:
                steps_lines.append(f"{step_num}. {step_text.strip()}")
            if steps_lines:
                case_data['steps'] = steps_lines
            else:
                steps_match = re.search(r'\*\*测试步骤\*\*：\n([\s\S]*?)(?=\n\*\*预期结果|$)', case_content, re.DOTALL)
                if steps_match:
                    steps_text = steps_match.group(1).strip()
                    case_data['steps'] = [s.strip() for s in steps_text.split('\n') if s.strip()]

            expected_lines = []
            exp_pattern = r'(?:^|\n)(\d+)[.、]\s*(.+)'
            exp_matches = re.findall(exp_pattern, case_content)
            for exp_num, exp_text in exp_matches:
                expected_lines.append(f"{exp_num}. {exp_text.strip()}")
            if expected_lines:
                case_data['expected_results'] = expected_lines
            else:
                exp_match = re.search(r'\*\*预期结果\*\*：\n([\s\S]*?)(?=\n\*\*实际结果|$)', case_content, re.DOTALL)
                if exp_match:
                    exp_text = exp_match.group(1).strip()
                    case_data['expected_results'] = [s.strip() for s in exp_text.split('\n') if s.strip()]

            self.test_cases.append(case_data)


class ExcelGenerator:
    """生成 Excel 格式的测试用例文档"""

    def __init__(self, parser, columns=None, template_path=None):
        """
        初始化 Excel 生成器

        Args:
            parser: 测试用例解析器
            columns: 列顺序列表，默认为 DEFAULT_COLUMNS
            template_path: 自定义模板路径（可选）
        """
        self.parser = parser
        self.template_path = template_path
        self.columns = columns or DEFAULT_COLUMNS
        self.wb = None
        self.ws = None
        self._init_workbook()

        if template_path:
            self._load_template_columns()

    def _init_workbook(self):
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

        self.wb = Workbook()
        self.ws = self.wb.active
        self.ws.title = "测试用例"

        self.header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        self.header_font = Font(bold=True, color="FFFFFF", size=11)
        self.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

    def _load_template_columns(self):
        """使用 TemplateParser 动态加载模板列定义"""
        from template_parser import TemplateParser

        parser = TemplateParser(self.template_path)
        field_mapping = parser.get_field_mapping()

        # 从 TemplateParser 获取字段映射
        self.columns = []
        self.column_to_json_key = {}

        for f in field_mapping:
            field_name = f["field_name"]
            json_key = f["json_key"]
            self.columns.append(field_name)
            self.column_to_json_key[field_name] = json_key

        # 保存 TemplateParser 供后续使用（如生成提示词）
        self.template_parser = parser

    def _set_cell_style(self, cell, value, is_header=False):
        from openpyxl.styles import Font, Alignment

        cell.value = value
        cell.border = self.border
        cell.alignment = Alignment(wrap_text=True, vertical='center')
        if is_header:
            cell.fill = self.header_fill
            cell.font = self.header_font
        else:
            cell.font = Font(size=10)

    def _format_steps(self, steps):
        if isinstance(steps, list):
            return '\n'.join(steps)
        return str(steps)

    def _format_expected_results(self, results):
        if isinstance(results, list):
            return '\n'.join(results)
        return str(results)

    def _get_case_value(self, case, column):
        """从用例数据中获取指定列的值"""
        # 如果使用模板，从模板的列映射中获取
        if self.template_path and hasattr(self, 'column_to_json_key'):
            key = self.column_to_json_key.get(column, column)
        else:
            # 默认列映射
            column_map = {
                '用例编号': 'id',
                '用例标题': 'name',
                '测试模块': 'module',
                '测试类型': 'test_type',
                '优先级': 'priority',
                '前置条件': 'precondition',
                '测试步骤': 'steps',
                '预期结果': 'expected_results',
                '测试数据': 'test_data',
                '备注': 'remarks',
                # 支持用户模板字段
                '操作': 'operation',
                '输入': 'input',
                '测试结果': 'test_result'
            }
            key = column_map.get(column, column)

        value = case.get(key, '')

        # 数组类型字段格式化
        if isinstance(value, (list, tuple)):
            value = '\n'.join(str(v) for v in value)
        elif key == 'test_data' and isinstance(value, dict):
            value = json.dumps(value, ensure_ascii=False)

        return value

    def generate(self, output_path):
        """生成 Excel 文件"""
        # 写入表头
        for col, header in enumerate(self.columns, 1):
            self._set_cell_style(self.ws.cell(row=1, column=col), header, is_header=True)

        # 写入数据
        for row, case in enumerate(self.parser.test_cases, 2):
            for col, header in enumerate(self.columns, 1):
                value = self._get_case_value(case, header)
                self._set_cell_style(self.ws.cell(row=row, column=col), value)

        # 设置列宽
        col_widths = {
            '用例编号': 15, '用例标题': 30, '测试模块': 20, '测试类型': 12,
            '优先级': 10, '前置条件': 30, '测试步骤': 40, '预期结果': 40,
            '测试数据': 25, '备注': 25
        }
        from openpyxl.utils import get_column_letter
        for i, col_name in enumerate(self.columns, 1):
            width = col_widths.get(col_name, 20)
            self.ws.column_dimensions[get_column_letter(i)].width = width

        self.wb.save(output_path)
        return output_path


class WordGenerator:
    """生成 Word 格式的测试用例文档"""

    def __init__(self, parser, template_path=None):
        self.parser = parser
        self.template_path = template_path
        self.doc = None

    def _init_document(self):
        from docx import Document
        from docx.oxml.ns import qn

        self.doc = Document()

    def _set_cell_font(self, cell, text, bold=False, size=10):
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(size)
                run.font.bold = bold
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _format_steps(self, steps):
        if isinstance(steps, list):
            return '\n'.join(steps)
        return str(steps)

    def _format_expected_results(self, results):
        if isinstance(results, list):
            return '\n'.join(results)
        return str(results)

    def generate(self, output_path):
        from docx.shared import Pt

        self._init_document()
        self.doc.add_heading('测试用例文档', 0)

        if self.parser.ticket_info:
            self.doc.add_heading('需求概述', level=1)
            table = self.doc.add_table(rows=len(self.parser.ticket_info), cols=2)
            table.style = 'Table Grid'
            for i, (key, value) in enumerate(self.parser.ticket_info.items()):
                self._set_cell_font(table.rows[i].cells[0], key, bold=True)
                self._set_cell_font(table.rows[i].cells[1], str(value))

        if self.parser.test_cases:
            self.doc.add_heading('测试用例', level=1)

            for case in self.parser.test_cases:
                self.doc.add_heading(f"{case['id']} {case['name']}", level=2)

                rows_data = [
                    ('用例编号', case['id']),
                    ('用例标题', case['name']),
                    ('测试模块', case.get('module', '')),
                    ('测试类型', case.get('test_type', '功能测试')),
                    ('优先级', case.get('priority', '中')),
                    ('前置条件', case['precondition'] if isinstance(case['precondition'], str) else '\n'.join(case['precondition']) if case['precondition'] else ''),
                    ('测试步骤', self._format_steps(case.get('steps', []))),
                    ('预期结果', self._format_expected_results(case.get('expected_results', []))),
                    ('测试数据', case.get('test_data', '')),
                    ('备注', case.get('remarks', ''))
                ]

                table = self.doc.add_table(rows=len(rows_data), cols=2)
                table.style = 'Table Grid'

                for i, (key, value) in enumerate(rows_data):
                    self._set_cell_font(table.rows[i].cells[0], key, bold=True)
                    self._set_cell_font(table.rows[i].cells[1], str(value))

                self.doc.add_paragraph()

        self.doc.save(output_path)
        return output_path


class MarkdownGenerator:
    """生成 Markdown 格式的测试用例文档"""

    def __init__(self, parser, template_path=None):
        self.parser = parser
        self.template_path = template_path

    def _format_steps(self, steps):
        if isinstance(steps, list):
            return '\n'.join(f'{i+1}. {s}' if not s.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) else s for i, s in enumerate(steps))
        return steps

    def _format_expected_results(self, results):
        if isinstance(results, list):
            return '\n'.join(f'{i+1}. {r}' if not r.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) else r for i, r in enumerate(results))
        return results

    def generate(self, output_path):
        lines = []
        lines.append("# 测试用例文档\n")

        if self.parser.ticket_info:
            lines.append("## 基本信息\n")
            lines.append("| 项目 | 内容 |")
            lines.append("|------|------|")
            for key, value in self.parser.ticket_info.items():
                lines.append(f"| {key} | {value} |")
            lines.append("")

        if self.parser.test_cases:
            lines.append("## 测试用例\n")
            for case in self.parser.test_cases:
                lines.append(f"### 测试用例编号：{case['id']} {case['name']}\n")
                lines.append(f"**测试模块**：{case.get('module', '')}")
                lines.append(f"**测试类型**：{case.get('test_type', '功能测试')}")
                lines.append(f"**优先级**：{case.get('priority', '中')}")

                precondition = case['precondition'] if isinstance(case['precondition'], str) else '\n'.join(case['precondition']) if case['precondition'] else ''
                if precondition:
                    lines.append(f"\n**前置条件**：\n{precondition}")

                lines.append(f"\n**测试步骤**：\n{self._format_steps(case.get('steps', []))}")
                lines.append(f"\n**预期结果**：\n{self._format_expected_results(case.get('expected_results', []))}")

                if case.get('test_data'):
                    lines.append(f"\n**测试数据**：\n{case['test_data']}")
                if case.get('remarks'):
                    lines.append(f"\n**备注**：{case.get('remarks')}")

                lines.append("\n---\n")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

        return output_path


def convert_to_docs(input_path, output_dir=None, template_path=None):
    """
    将 JSON 或 Markdown 测试用例转换为 Markdown、Word 和 Excel 格式

    Args:
        input_path: 输入文件路径 (.json 或 .md)
        output_dir: 输出目录（默认与输入文件同目录）
        template_path: 自定义模板文件路径（可选）
                     如果提供，则根据模板格式生成对应文档
                     如果不提供，则生成 MD、DOCX、XLSX 默认格式

    Returns:
        生成的文档路径列表
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"文件不存在: {input_path}")

    if output_dir is None:
        output_dir = os.path.dirname(input_path)

    ext = os.path.splitext(input_path)[1].lower()

    if ext == '.json':
        parser = JSONTestCaseParser(input_path)
        print(f"解析到 ticket_info: {len(parser.ticket_info)} 条")
        print(f"解析到 test_cases: {len(parser.test_cases)} 条")
    elif ext == '.md':
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        parser = MarkdownTestCaseParser(md_content)
        print(f"解析到 ticket_info: {len(parser.ticket_info)} 条")
        print(f"解析到 test_cases: {len(parser.test_cases)} 条")
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .json 和 .md")

    base_name = os.path.splitext(os.path.basename(input_path))[0]
    generated_files = []

    # 如果提供了模板，根据模板格式生成对应文档
    if template_path:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        template_ext = os.path.splitext(template_path)[1].lower()

        if template_ext == '.xlsx':
            # Excel 模板 → 生成 XLSX
            xlsx_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.xlsx"))
            ExcelGenerator(parser, template_path=template_path).generate(xlsx_path)
            print(f"Excel 文档已生成（使用模板）: {xlsx_path}")
            generated_files = [xlsx_path]

        elif template_ext == '.docx':
            # Word 模板 → 生成 DOCX（使用模板格式）
            docx_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.docx"))
            WordGenerator(parser, template_path=template_path).generate(docx_path)
            print(f"Word 文档已生成（使用模板）: {docx_path}")
            generated_files = [docx_path]

        elif template_ext == '.md':
            # Markdown 模板 → 生成 MD（使用模板格式）
            md_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.md"))
            MarkdownGenerator(parser, template_path=template_path).generate(md_path)
            print(f"Markdown 文档已生成（使用模板）: {md_path}")
            generated_files = [md_path]

        else:
            raise ValueError(f"不支持的模板格式: {template_ext}，仅支持 .xlsx、.docx、.md")
    else:
        # 没有提供模板，生成 MD、DOCX、XLSX 默认格式
        md_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.md"))
        MarkdownGenerator(parser).generate(md_path)
        print(f"Markdown 文档已生成: {md_path}")
        generated_files.append(md_path)

        docx_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.docx"))
        WordGenerator(parser).generate(docx_path)
        print(f"Word 文档已生成: {docx_path}")
        generated_files.append(docx_path)

        xlsx_path = os.path.normpath(os.path.join(output_dir, f"{base_name}.xlsx"))
        ExcelGenerator(parser).generate(xlsx_path)
        print(f"Excel 文档已生成: {xlsx_path}")
        generated_files.append(xlsx_path)

    return generated_files


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='将 JSON 或 Markdown 测试用例转换为 Markdown、Word 和 Excel 格式')
    parser.add_argument('input_file', help='JSON 或 Markdown 测试用例文件路径')
    parser.add_argument('-o', '--output', help='输出目录（默认与输入文件同目录）')
    parser.add_argument('-t', '--template', help='自定义 Excel 模板文件路径（可选）')

    args = parser.parse_args()

    try:
        convert_to_docs(args.input_file, args.output, args.template)
        print("\n转换完成！")
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
        exit(1)